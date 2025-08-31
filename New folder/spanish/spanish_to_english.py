import os
import threading
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from transformers import MarianMTModel, MarianTokenizer, pipeline
from docx import Document
import torch

# ---------------- Configuration ----------------
MODEL_NAME = "Helsinki-NLP/opus-mt-es-en"   # Spanish -> English
MODEL_DIR = r"C:\Users\intel\models\opus-mt-es-en"

translator = None
device_index = 0 if torch.cuda.is_available() else -1
device_str = "cuda" if torch.cuda.is_available() else "cpu"

# ---------------- Model loading ----------------
def load_model_async(root_window, status_label, translate_button):
    def _load():
        global translator
        try:
            if not os.path.isdir(MODEL_DIR) or not os.listdir(MODEL_DIR):
                root_window.after(0, lambda: status_label.config(text=f"Downloading model {MODEL_NAME} ..."))
                model = MarianMTModel.from_pretrained(MODEL_NAME)
                tokenizer = MarianTokenizer.from_pretrained(MODEL_NAME)
                os.makedirs(MODEL_DIR, exist_ok=True)
                model.save_pretrained(MODEL_DIR)
                tokenizer.save_pretrained(MODEL_DIR)

            tokenizer = MarianTokenizer.from_pretrained(MODEL_DIR)
            model = MarianMTModel.from_pretrained(MODEL_DIR)
            model.to(device_str)
            translator = pipeline("translation", model=model, tokenizer=tokenizer, device=device_index)
            root_window.after(0, lambda: status_label.config(text=f"Model loaded ({device_str}). Ready to translate."))
            root_window.after(0, lambda: translate_button.config(state="normal"))
        except Exception as ex:
            root_window.after(0, lambda e=ex: messagebox.showerror("Model load error", f"Error loading model:\n\n{e}"))
            root_window.after(0, lambda: status_label.config(text="Model load failed."))

    threading.Thread(target=_load, daemon=True).start()

# ---------------- Translation ----------------
def translate_paragraphs(text, progress_var, progress_bar, status_label):
    if translator is None:
        raise RuntimeError("Translator is not loaded.")

    paragraphs = text.split("\n\n")
    total = len(paragraphs) if paragraphs else 1
    out_paragraphs = []

    for idx, para in enumerate(paragraphs, start=1):
        if para.strip():
            try:
                result = translator(para, max_length=512)
                translated = result[0]["translation_text"]
            except Exception as e:
                translated = f"[ERROR translating paragraph: {e}]"
        else:
            translated = ""
        out_paragraphs.append(translated)

        progress_var.set(int(idx / total * 100))
        progress_bar.update_idletasks()
        status_label.config(text=f"Translating... {int(idx / total * 100)}%")

    return "\n\n".join(out_paragraphs)

# ---------------- GUI Handlers ----------------
def on_open_docx(input_widget):
    fp = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
    if not fp:
        return
    try:
        doc = Document(fp)
        input_widget.delete("1.0", tk.END)
        text = "\n\n".join([p.text for p in doc.paragraphs])
        input_widget.insert("1.0", text)
    except Exception as e:
        messagebox.showerror("Open error", f"Failed to open .docx:\n\n{e}")

def on_save_docx(output_widget):
    fp = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
    if not fp:
        return
    try:
        doc = Document()
        full_text = output_widget.get("1.0", tk.END).rstrip()
        if full_text:
            for para in full_text.split("\n\n"):
                doc.add_paragraph(para)
        doc.save(fp)
        messagebox.showinfo("Saved", f"Saved to:\n{fp}")
    except Exception as e:
        messagebox.showerror("Save error", f"Failed to save .docx:\n\n{e}")

def on_clear_input(input_widget):
    input_widget.delete("1.0", tk.END)

def on_translate(input_widget, output_widget, progress_var, progress_bar, status_label):
    text = input_widget.get("1.0", tk.END).strip()
    if not text:
        messagebox.showwarning("Empty input", "Please enter text to translate.")
        return

    translate_btn.config(state="disabled")
    output_widget.config(state="normal")
    output_widget.delete("1.0", tk.END)
    progress_var.set(0)
    status_label.config(text="Starting translation...")

    def _task():
        try:
            translated = translate_paragraphs(text, progress_var, progress_bar, status_label)
            root.after(0, lambda: output_widget.insert("1.0", translated))
            root.after(0, lambda: status_label.config(text="Translation complete."))
            root.after(0, lambda: progress_var.set(100))
        except Exception as e:
            root.after(0, lambda: messagebox.showerror("Translation error", f"{e}"))
            root.after(0, lambda: status_label.config(text="Translation failed."))
        finally:
            root.after(0, lambda: translate_btn.config(state="normal"))
            root.after(0, lambda: output_widget.config(state="disabled"))

    threading.Thread(target=_task, daemon=True).start()

# ---------------- GUI Build ----------------
root = tk.Tk()
root.title("Spanish → English Translator")
root.geometry("950x650")
root.configure(bg="#1a1a2e")

# Header
header = tk.Label(root, text="🇪🇸 Spanish → English Translator",
                  font=("Helvetica", 24, "bold"),
                  bg="#162447", fg="#e43f5a", pady=12)
header.pack(fill=tk.X)

# Input frame
input_frame = tk.Frame(root, bg="#1f4068")
input_frame.pack(fill=tk.X, padx=28, pady=(18, 8))
tk.Label(input_frame, text="Enter Spanish text or paste content:",
         bg="#1f4068", fg="#e0e0e0", font=("Helvetica", 12, "bold")).pack(anchor="w", padx=8, pady=(6, 0))
input_text = tk.Text(input_frame, height=12, wrap="word",
                     bg="#e0e0e0", fg="#1f4068", font=("Helvetica", 12), padx=10, pady=10, relief="flat")
input_text.pack(fill=tk.BOTH, padx=8, pady=8)

# Buttons row
btn_frame = tk.Frame(root, bg="#1a1a2e")
btn_frame.pack(fill=tk.X, padx=28, pady=(0, 14))

translate_btn = tk.Button(btn_frame, text="Translate", bg="#e43f5a", fg="white",
                          font=("Helvetica", 12, "bold"), padx=20, pady=8,
                          state="disabled",
                          command=lambda: on_translate(input_text, output_text, progress_var, progress_bar, status_label))
translate_btn.pack(side="left", padx=8)

clear_btn = tk.Button(btn_frame, text="Clear Input", bg="#0f3460", fg="white",
                      font=("Helvetica", 12, "bold"), padx=20, pady=8,
                      command=lambda: on_clear_input(input_text))
clear_btn.pack(side="left", padx=8)

open_btn = tk.Button(btn_frame, text="Open .docx", bg="#0f3460", fg="white",
                     font=("Helvetica", 12, "bold"), padx=20, pady=8,
                     command=lambda: on_open_docx(input_text))
open_btn.pack(side="left", padx=8)

save_btn = tk.Button(btn_frame, text="Save Translation", bg="#0f3460", fg="white",
                     font=("Helvetica", 12, "bold"), padx=20, pady=8,
                     command=lambda: on_save_docx(output_text))
save_btn.pack(side="right", padx=8)

# Output frame
output_frame = tk.Frame(root, bg="#1f4068")
output_frame.pack(fill=tk.BOTH, expand=True, padx=28, pady=(0, 12))
tk.Label(output_frame, text="English Translation:",
         bg="#1f4068", fg="#e0e0e0", font=("Helvetica", 12, "bold")).pack(anchor="w", padx=8, pady=(6, 0))
output_text = tk.Text(output_frame, height=12, wrap="word",
                      bg="#e0e0e0", fg="#1f4068", font=("Helvetica", 12), padx=10, pady=10, relief="flat")
output_text.pack(fill=tk.BOTH, expand=True, padx=8, pady=8)
output_text.insert("1.0", "Translation will appear here...")
output_text.config(state="disabled")

# Progress bar & status
progress_var = tk.IntVar(value=0)
progress_bar = ttk.Progressbar(root, orient="horizontal", length=700, mode="determinate", variable=progress_var)
progress_bar.pack(pady=(4, 6))

status_label = tk.Label(root, text="Loading model...", bg="#1a1a2e", fg="#00adb5", font=("Helvetica", 10))
status_label.pack(side="bottom", fill=tk.X, pady=6)

# Load model async
load_model_async(root, status_label, translate_btn)

root.mainloop()
